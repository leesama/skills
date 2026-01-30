import argparse
import json
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font_style(paragraph, font_name="SimHei", font_size=12):
    """设置段落字体样式为简体中文"""
    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Inches(font_size / 72)


def create_combined_task_table(doc, this_tasks, next_tasks, period_type="周"):
    """创建合并的任务表格"""
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    title_row = table.rows[0]
    title_row.cells[0].merge(title_row.cells[3])
    title_row.cells[0].text = f"本{period_type}任务"
    title_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_style(title_row.cells[0].paragraphs[0])

    header_row = table.add_row()
    headers = ["任务内容", "完成标准", "完成状态", "备注"]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_style(header_row.cells[i].paragraphs[0])

    for task in this_tasks:
        row_cells = table.add_row().cells
        task_data = task[:4] if len(task) >= 4 else task + [""] * (4 - len(task))
        for i, content in enumerate(task_data):
            row_cells[i].text = content
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font_style(row_cells[i].paragraphs[0])

    title_row = table.add_row()
    title_row.cells[0].merge(title_row.cells[3])
    title_row.cells[0].text = f"下{period_type}任务"
    title_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_style(title_row.cells[0].paragraphs[0])

    header_row = table.add_row()
    next_headers = ["任务内容", "完成标准", "备注", ""]
    for i, header in enumerate(next_headers):
        header_row.cells[i].text = header
        header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_style(header_row.cells[i].paragraphs[0])

    for task in next_tasks:
        row_cells = table.add_row().cells
        extended_task = task + [""] * (4 - len(task))
        for i, content in enumerate(extended_task):
            row_cells[i].text = content
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font_style(row_cells[i].paragraphs[0])

    return table


def _period_type_from_report(report_type):
    if not report_type:
        return "周"
    return "月" if "月" in report_type else "周"


def render_word(data, output_file=None):
    period = data.get("period") or {}
    report_type = data.get("report_type", "")
    period_type = _period_type_from_report(report_type)
    start_date = period.get("start_date", "")
    end_date = period.get("end_date", "")
    total_commits = (data.get("statistics") or {}).get("total_commits", "")

    tasks = data.get("tasks") or []
    this_tasks = [
        [
            t.get("content", ""),
            t.get("completion_standard", ""),
            t.get("status", ""),
            t.get("notes", ""),
        ]
        for t in tasks
    ]

    next_tasks = data.get("next_tasks") or [["", "", ""]]

    doc = Document()
    heading = doc.add_heading(f"本{period_type}工作{period_type}报", level=1)
    set_font_style(heading)

    time_para = doc.add_paragraph(f"统计时间：{start_date} 至 {end_date}")
    set_font_style(time_para)

    count_para = doc.add_paragraph(f"本{period_type}提交总数：{total_commits} 条")
    set_font_style(count_para)

    create_combined_task_table(doc, this_tasks, next_tasks, period_type)

    if not output_file:
        output_file = f"本{period_type}工作{period_type}报_{end_date}.docx"
    doc.save(output_file)
    return output_file


def main():
    parser = argparse.ArgumentParser(description="将优化后的周/月报 JSON 数据渲染成 Word")
    parser.add_argument("-i", "--input", required=True, help="优化后的 JSON 文件路径")
    parser.add_argument("-o", "--output", default="", help="输出的 Word 文件路径（可选）")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    output_file = render_word(data, args.output or None)
    print(f"✅ Word 已生成：{output_file}")


if __name__ == "__main__":
    main()
